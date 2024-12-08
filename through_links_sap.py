import time
from typing import List

from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.chrome.webdriver import WebDriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import StaleElementReferenceException, NoSuchElementException, NoSuchWindowException
from selenium.webdriver.remote.shadowroot import ShadowRoot

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_shadow_root(driver: WebDriver, element: WebElement) -> ShadowRoot:
    return driver.execute_script('return arguments[0].shadowRoot', element)

# Возможности и решения
class SolutionCapability:
    name: str
    description_short: str
    #description_long: str
    #required_software_products: List[str]

    def __init__(self, **kwargs):
        self.name = ""
        self.required_software_products = []

# Деловые возможности
class BusinessCapability:
    name: str
    desc: str
    solution_capabilities: List[SolutionCapability]

    def __init__(self, **kwargs):
        self.name = ""
        self.solution_capabilities = []

# Область бизнеса
class BusinessArea:
    name: str
    business_capabilities: List[BusinessCapability]

    def __init__(self, **kwargs):
        self.name = ""
        self.business_capabilities = []

# Ветка бизнеса
class LineOfBusiness:
    name: str
    business_areas: List[BusinessArea]

    def __init__(self, **kwargs):
        self.name = ""
        self.business_areas = []

# Сведения об индустрии
class Industry:
    lines_of_business: List[LineOfBusiness]
    lines_of_technologies: List[LineOfBusiness]

    def __init__(self, **kwargs):
        self.lines_of_business = []
        self.lines_of_technologies = []

# Индустрия -> Ветки бизнеса -> Область бизнеса -> Возможности и решения -> информация
# Healthcare -> Asset Management -> Asset Acquisition and Building -> Project Financials Control -> информация

# приходим сюда с классом <a> sl-sub-element
def parse_solution_capability_page(driver: WebDriver, area: WebElement) -> SolutionCapability:
    print("\t\t\t\tparse_solution_capability_page: start")
    res = SolutionCapability()

    href = area.get_attribute('href')
    original_window = driver.current_window_handle
    driver.execute_script("window.open(arguments[0])", href)
    #time.sleep(1)

    new_window = driver.window_handles[-1]
    driver.switch_to.window(new_window)

    # новая страница
    # classes:
    # <solution-library-web-object> sl-layout-content-container -- содержит внутри себя shadow root, который нужно извлекать
    # 
    # <div> sl-name -- контейнер с названием solution_capability
    # <div> sl-svp -- контейнер с коротким описанием
    # <div> sl-parent-description -- контейнер с длинным описанием 1
    # <div> sl-description -- контейнер с длинным описанием 2
    # 
    # <div> sl-section-spacing 
    #   <div> sl-attribute-name -- контейнер с необходимым продуктом (список, но мб не нужно)

    # ждем, пока не подтянется контейнер с shadow root
    WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH, '//solution-library-web-object[@class="sl-layout-content-container"]')))
    #extra wait to make all the quotes loaded
    time.sleep(1)
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")

    # достаем контейнер с shadow root
    shadow_element_container = driver.find_element(By.XPATH, '//solution-library-web-object[@class="sl-layout-content-container"]')
    shadow_root = get_shadow_root(driver, shadow_element_container)
    # содержимое shadow root
    internal_data = shadow_root.find_element(By.CSS_SELECTOR, 'div')

    sl_name_element = internal_data.find_element(By.XPATH, './/div[@class="sl-name"]')
    res.name = sl_name_element.text

    sl_svp_element = internal_data.find_element(By.XPATH, './/div[@class="sl-svp"]')
    res.description_short = sl_svp_element.text

    #sl_parent_description_element = internal_data.find_element(By.XPATH, './/div[@class="sl-parent-description"]')
    #res.description_long = sl_parent_description_element.text

    #sl_description_element = internal_data.find_element(By.XPATH, './/div[@class="sl-description"]')
    #res.description_long += '. ' + sl_description_element.text

    #Close the tab or window
    driver.close()

    #Switch back to the old tab or window
    driver.switch_to.window(original_window)
    return res

# приходим с классом sl-two-column-item
def parse_business_capability(driver: WebDriver, capability: WebElement) -> BusinessCapability:
    print("\t\t\tparse_business_capability: start")
    res = BusinessCapability()

    # classes:
    # <solution-library-web-bcm-object> sl-layout-content-container -- содержит внутри себя shadow root, который нужно извлекать
    # 
    # <div> sl-name -- контейнер с названием business_area
    # <div> sl-two-column-container -- контейнер со всеми business capabilities.
    #   <div> sl-two-column-item -- контейнер конкретной business capability
    #       <div> sl-content-name -- название конкретной business capability
    #       <div> sl-content-description -- описание конкретной business capability
    #       <a> sl-sub-element -- содержит href с названием solution capability внутри 

    business_capability_name_element = capability.find_element(By.XPATH, './/div[@class="sl-content-name"]')
    res.name = business_capability_name_element.text

    business_capability_desc_element = capability.find_element(By.XPATH, './/div[@class="sl-content-description"]')
    res.desc = business_capability_desc_element.text

    return res

def parse_business_area_page(driver: WebDriver, area: WebElement) -> BusinessArea:
    print("\t\tparse_business_area_page: start")

    res = BusinessArea()

    href = area.get_attribute('href')
    original_window = driver.current_window_handle
    driver.execute_script("window.open(arguments[0])", href)
    #time.sleep(1)

    new_window = driver.window_handles[-1]
    driver.switch_to.window(new_window)

    # новая страница
    # classes:
    # <solution-library-web-bcm-object> sl-layout-content-container -- содержит внутри себя shadow root, который нужно извлекать
    # 
    # <div> sl-name -- контейнер с названием business_area
    # <div> sl-two-column-item -- контейнер со всеми business capabilities.
    #   <div> sl-two-column-item -- контейнер конкретной business capability
    #       <div> sl-content-name -- название конкретной business capability
    #       <div> sl-content-description -- описание конкретной business capability
    #       <a> sl-sub-element -- содержит href с названием solution capability внутри 

    # ждем, пока не подтянется контейнер с shadow root
    WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH, '//solution-library-web-bcm-object[@class="sl-layout-content-container"]')))
    #extra wait to make all the quotes loaded
    time.sleep(1)
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")

    # достаем контейнер с shadow root
    shadow_element_container = driver.find_element(By.XPATH, '//solution-library-web-bcm-object[@class="sl-layout-content-container"]')
    shadow_root = get_shadow_root(driver, shadow_element_container)
    # содержимое shadow root
    internal_data = shadow_root.find_element(By.CSS_SELECTOR, 'div')

    sl_name_element = internal_data.find_element(By.XPATH, './/div[@class="sl-name"]')
    res.name = sl_name_element.text

    business_capabilities = internal_data.find_elements(By.XPATH, './/div[@class="sl-two-column-item"]')
    for index, val in enumerate(business_capabilities):
        print("\t\t\tindex = %d business_capability" % index)
        try:
            #get the qutoes again after getting back to the initial page in the loop
            business_capabilities = internal_data.find_elements(By.XPATH, './/div[@class="sl-two-column-item"]')
            business_capability = parse_business_capability(driver, business_capabilities[index])

            res.business_capabilities.append(business_capability)

            #WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH, './/div[@class="sl-sub-element"]')))
            #print("HEY")
            #time.sleep(0.1)
        except StaleElementReferenceException:  
            pass

    #Close the tab or window
    driver.close()

    #Switch back to the old tab or window
    driver.switch_to.window(original_window)

    return res

# приходим с классом sl-vm-object-two-column-item
def parse_line_of_business(driver: WebDriver, line: WebElement) -> LineOfBusiness:
    print("\tparse_line_of_business: start")

    # classes:
    # <solution-library-web-bcm> sl-layout-content-container -- содержит внутри себя shadow root, который нужно извлекать
    # 
    # <div> sl-vm-object-bp-container -- контейнер со всеми line-items. Таких контейнера два! (первый - line items, второй - technologies)
    #   <div> sl-vm-object-two-column-item -- контейнер конкретной line-item
    #       <div> sl-content-name -- название конкретной line-item
    #       <a> sl-sub-element -- содержит href с названием business area внутри 
    #           <img> sl-sub-element-icon -- содержит название business area

    res = LineOfBusiness()

    line_of_business_name_element = line.find_element(By.XPATH, './/div[@class="sl-content-name"]')
    res.name = line_of_business_name_element.text

    business_areas = line.find_elements(By.XPATH, './/a[@class="sl-sub-element"]')
    for index, val in enumerate(business_areas):
        print("\tindex = %d parse_line_of_business" % index)
        try:
            #get the qutoes again after getting back to the initial page in the loop
            business_areas = line.find_elements(By.XPATH, './/a[@class="sl-sub-element"]')
            business_area = parse_business_area_page(driver, business_areas[index])

            res.business_areas.append(business_area)

            #WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH, './/div[@class="sl-sub-element"]')))
            #print("HEY")
            #time.sleep(0.1)
        except StaleElementReferenceException:  
            pass
    
    return res

def parse_industry_page(driver: WebDriver, industry_page_url: str) -> Industry:
    print("parse_industry_page: start")
    # classes:
    # <solution-library-web-bcm> sl-layout-content-container -- содержит внутри себя shadow root, который нужно извлекать
    # 
    # <div> sl-vm-object-bp-container -- контейнер со всеми line-items. Таких контейнера два! (первый - line items, второй - technologies)
    #   <div> sl-vm-object-two-column-item -- контейнер конкретной line-item
    #       <div> sl-content-name -- название конкретной line-item
    #       <a> sl-sub-element -- содержит href с названием business area внутри 
    #           <img> sl-sub-element-icon -- содержит название business area

    # результат
    res = Industry()

    driver.get(industry_page_url)

    # ждем, пока не подтянется контейнер с shadow root
    WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH, '//solution-library-web-bcm[@class="sl-layout-content-container"]')))
    #extra wait to make all the quotes loaded
    time.sleep(1)
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")

    # достаем контейнер с shadow root
    shadow_element_container = driver.find_element(By.XPATH, '//solution-library-web-bcm[@class="sl-layout-content-container"]')
    shadow_root = get_shadow_root(driver, shadow_element_container)
    # содержимое shadow root
    internal_datas = shadow_root.find_elements(By.CLASS_NAME, "sl-vm-object-bp-container")

    lines_of_business = internal_datas[0].find_elements(By.XPATH, './/div[@class="sl-vm-object-two-column-item"]')
    print("number of lines of business = %d" % len(lines_of_business))

    time.sleep(0.1)
    for index, val in enumerate(lines_of_business):
        print("index = %d parse_industry_page" % index)
        try:
            #get the qutoes again after getting back to the initial page in the loop
            lines_of_business = internal_datas[0].find_elements(By.XPATH, './/div[@class="sl-vm-object-two-column-item"]')

            start_time = time.time()

            line_of_business = parse_line_of_business(driver, lines_of_business[index])
            res.lines_of_business.append(line_of_business)

            end_time = time.time()
            elapsed_time = end_time - start_time
            print('got [%d] business line info in %f seconds' % (index, elapsed_time))

            #WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH, './/div[@class="sl-vm-object-two-column-item"]')))
            #time.sleep(0.1)
        except StaleElementReferenceException:  
            pass

    lines_of_technologies = internal_datas[1].find_elements(By.XPATH, './/div[@class="sl-vm-object-two-column-item"]')
    print("number of lines of technologies = %d" % len(lines_of_technologies))

    time.sleep(0.1)
    for index, val in enumerate(lines_of_technologies):
        print("index = %d parse_industry_page" % index)
        try:
            #get the qutoes again after getting back to the initial page in the loop
            lines_of_technologies = internal_datas[1].find_elements(By.XPATH, './/div[@class="sl-vm-object-two-column-item"]')

            start_time = time.time()

            line_of_business = parse_line_of_business(driver, lines_of_technologies[index])
            res.lines_of_technologies.append(line_of_business)

            end_time = time.time()
            elapsed_time = end_time - start_time
            print('got [%d] technology line info in %f seconds' % (index, elapsed_time))

            #WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH, './/div[@class="sl-vm-object-two-column-item"]')))
            #time.sleep(0.1)
        except StaleElementReferenceException:  
            pass

    return res

def print_to_word(industry: Industry):
    document = Document()

    normal_style = document.styles['Normal']
    normal_style_font = normal_style.font
    normal_style_font.name = 'Times New Roman'

    for i, line in enumerate(industry.lines_of_business):
        document.add_heading(line.name, level=1)
        for j, business_area in enumerate(line.business_areas):
            document.add_heading(business_area.name, level=2)
            for k, business_capability in enumerate(business_area.business_capabilities):
                document.add_heading(business_capability.name, level=3)
                paragraph = document.add_paragraph().add_run(business_capability.desc)
                paragraph.font.size = Pt(12)
    
    for i, line in enumerate(industry.lines_of_technologies):
        document.add_heading(line.name, level=1)
        for j, business_area in enumerate(line.business_areas):
            document.add_heading(business_area.name, level=2)
            for k, business_capability in enumerate(business_area.business_capabilities):
                document.add_heading(business_capability.name, level=3)
                paragraph = document.add_paragraph().add_run(business_capability.desc)
                paragraph.font.size = Pt(12)

    document.add_page_break()

    document.save("health.docx")

    return

def main():
    print("Hello, World!")

    driver = webdriver.Chrome()
    industry_name = "HEALTH"
    industry_page_url = "https://solutionportfolio.net.sap/bcm/industry/" + industry_name
    start_time = time.time()
    industry = parse_industry_page(driver, industry_page_url)
    driver.quit()
    end_time = time.time()
    elapsed_time = end_time - start_time
    print('got industry info in: ', elapsed_time)
    
    start_time = time.time()
    print_to_word(industry)
    end_time = time.time()
    elapsed_time = end_time - start_time
    print('printed to word in: ', elapsed_time)


if __name__ == "__main__":
    main()
